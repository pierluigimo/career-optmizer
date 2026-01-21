import streamlit as st
import yaml
import re
import os
import urllib.parse
from docx import Document
from docx.shared import Pt
import io
from datetime import datetime

# --- CONFIGURAZIONE PAGINA ---
st.set_page_config(page_title="Strategic Career Optimizer", page_icon="🚀", layout="wide")

# --- DATI MASTER (ESTRATTI DAI TUOI PDF) ---
# Incorporati qui per evitare problemi di lettura file sul cloud
PROFILE_YAML = """
meta:
  user_id: "pierluigi_monaco"
  role_type: "Strategic Finance Executive"
  contact_info:
    phone: "+39 348 8513668"
    email: "pierluigimonaco2@gmail.com"
    location: "Milano"
    linkedin: "linkedin.com/in/pierluigimonaco"

en:
  summary: "Strategic Finance Executive with 20+ years of global experience leading transformation, financial governance and operational optimization. Delivered 15% OPEX reduction through $9M rebate harmonization and stabilized operations in volatile markets. Influenced $6B+ in investments and partnered with C-level leaders."
  core_projects:
    - title: "Subscription Model Transformation"
      company: "OpenText"
      impact: "Generated $6.9M new subscription business & reduced churn from 20% to 10%."
    - title: "Global Financing Policy"
      company: "Micro Focus"
      impact: "Designed policy for $6B worldwide operations."
  capability_themes:
    - id: "strategic_finance"
      name: "Strategic Finance & Executive Partnering"
      keywords: ["C-Level", "Stakeholder", "Strategy", "Governance"]
      evidence:
        - "Managed $205M revenue and $95M cost portfolio across EMEA, NA, APAC."
        - "Influenced $6B+ global investment decisions."
        - "Supported 425+ FTEs globally."
    - id: "transformation"
      name: "Cloud / Subscription Transformation"
      keywords: ["SaaS", "Subscription", "Recurring Revenue", "Cloud"]
      evidence:
        - "Generated $6.9M new subscription business."
        - "Reduced contract loss rate from 20% to 10%."
        - "Transitioned financial models impacting $58M revenue base."
    - id: "cost_optimization"
      name: "Cost Optimization & Operational Excellence"
      keywords: ["Cost Reduction", "Efficiency", "OPEX", "Shared Services"]
      evidence:
        - "Achieved 15% OPEX reduction via process harmonization."
        - "Delivered 10% operational cost reduction across 3 regions."
  education:
    - "CFO Program (Ongoing) - Business School Il Sole 24 Ore (2025)"
    - "Master in Finance, Administration and Control - Business School Il Sole 24 Ore"
    - "Bachelor's Degree in Economics - Università Cattolica del Sacro Cuore"
  skills:
    hard: ["Financial Modeling", "IFRS/US GAAP", "SaaS Metrics (ARR/MRR)", "Risk Management"]
    tools: ["SAP", "OneStream", "Hyperion", "Essbase", "Excel (Advanced)"]

it:
  summary: "Strategic Finance Executive con oltre 20 anni di esperienza globale nella guida di progetti di trasformazione, governance finanziaria e ottimizzazione operativa. Ho conseguito una riduzione del 15% dei costi operativi (OPEX) e supportato oltre 4 milioni di dollari in nuovi ricavi cloud."
  core_projects:
    - title: "Trasformazione Modello Subscription"
      company: "OpenText"
      impact: "Generato $6.9M di nuovo business e ridotto il tasso di perdita dal 20% al 10%."
    - title: "Policy Globale di Finanziamento"
      company: "Micro Focus"
      impact: "Policy applicata su operazioni da oltre $6B a livello mondiale."
  capability_themes:
    - id: "strategic_finance"
      name: "Finanza Strategica & Executive Partnering"
      keywords: ["Direzione", "Strategia", "Stakeholder", "Governance"]
      evidence:
        - "Gestione portafoglio ricavi $205M e costi $95M su EMEA, NA, APAC."
        - "Advisor strategico per decisioni di investimento globali da oltre $6B."
    - id: "transformation"
      name: "Trasformazione Cloud / Subscription"
      keywords: ["SaaS", "Subscription", "Cloud", "Ricavi Ricorrenti"]
      evidence:
        - "Generato $6.9M in nuovo business subscription."
        - "Ridotto il tasso di perdita contrattuale dal 20% al 10%."
    - id: "cost_optimization"
      name: "Eccellenza Operativa"
      keywords: ["Riduzione Costi", "Efficienza", "OPEX"]
      evidence:
        - "Riduzione del 15% dei costi operativi (OPEX) tramite armonizzazione processi."
        - "-10% sui costi operativi in tre regioni."
  education:
    - "CFO Program (In corso) - Business School Il Sole 24 Ore (2025)"
    - "Master in Finanza, Amministrazione e Controllo - Business School Il Sole 24 Ore"
    - "Laurea in Economia e Gestione Aziendale - Università Cattolica del Sacro Cuore"
  skills:
    hard: ["Modellazione Finanziaria", "IFRS/US GAAP", "Metriche SaaS", "Gestione Rischio"]
    tools: ["SAP", "OneStream", "Hyperion", "Essbase", "Excel (Avanzato)"]
"""

# --- LOGICA DI BUSINESS ---
def detect_language(text):
    text = text.lower()
    score_en = sum(text.count(w) for w in [' the ', ' and ', ' to ', ' requirements ', ' we are '])
    score_it = sum(text.count(w) for w in [' il ', ' e ', ' per ', ' requisiti ', ' siamo '])
    return 'it' if score_it > score_en else 'en'

def extract_info(jd_text, lang):
    # Company Name (euristica semplice)
    match = re.search(r'(?:About|Join|working at|presso|azienda)\s+([A-Z][a-zA-Z0-9\s\&]+?)(?:\s|\.|,)', jd_text)
    company = match.group(1).strip() if match and len(match.group(1)) < 30 else "Company"
    
    # Role Title
    role = "Finance Executive"
    if "VP" in jd_text: role = "VP Finance"
    if "Director" in jd_text: role = "Finance Director"
    if "Head" in jd_text: role = "Head of Finance"
    
    # Characteristic (Key Value)
    keywords = ['innovation', 'sustainability', 'leadership', 'growth', 'transformation'] if lang == 'en' else ['innovazione', 'sostenibilità', 'leadership', 'crescita', 'trasformazione']
    char = "mission" if lang == 'en' else "missione"
    for k in keywords:
        if k in jd_text.lower():
            char = k
            break
            
    return company, role, char

def create_docs(profile, lang, company, role, char):
    content = profile[lang]
    meta = profile['meta']
    
    # Selezione Temi (MVP: prende i primi 2, in futuro puoi aggiungere logica scoring)
    theme1 = content['capability_themes'][0]
    theme2 = content['capability_themes'][1]
    ev1 = theme1['evidence'][0]
    ev2 = theme2['evidence'][0]
    
    # --- 1. COVER LETTER (Template PDF "Risposta Annuncio Azienda") ---
    doc_cl = Document()
    style = doc_cl.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Intestazione
    doc_cl.add_paragraph(f"{meta['user_id'].replace('_',' ').title()} | {meta['contact_info']['email']} | {meta['contact_info']['phone']}")
    doc_cl.add_paragraph(f"\nSubject: Application for {role} - {company}" if lang == 'en' else f"\nOggetto: Candidatura per {role} - {company}")
    
    if lang == 'en':
        body = f"""
Dear Hiring Team,

I am writing to express my deep interest in the {role} position at {company}.

My admiration for {company} is grounded not only in its leadership role, but also in the company's commitment to {char}. I firmly believe that my professional career reflects these same values, creating a natural alignment between your organization's goals and mine.

In my current role as Strategic Finance Executive, I have developed a deep understanding of {theme1['name']}. My dedication and results, specifically {ev1.lower()}, demonstrate not only my capability in {theme1['keywords'][0]} but also an alignment with the results-oriented approach that {company} embodies.

Furthermore, my experience in {theme2['name']} ({ev2.lower()}) would complement your current team, contributing to guiding the company toward new milestones.

Thank you for considering my application. I am available for an interview to discuss how I can contribute to your success.

Sincerely,

Pierluigi Monaco
"""
    else:
        body = f"""
Gentile Team di Selezione,

Scrivo questa lettera per esprimere il mio profondo interesse per la posizione di {role} presso {company}.

La mia ammirazione per {company} non è solo radicata nel suo ruolo di leader, ma anche nell'impegno dell'azienda verso {char}. Credo fermamente che la mia carriera professionale rifletta questi stessi valori, creando un'armonia naturale tra gli obiettivi della vostra realtà ed i miei.

Nel mio attuale ruolo di Strategic Finance Executive, ho sviluppato una profonda comprensione di {theme1['name']}. Il mio impegno e i risultati ottenuti, specialmente {ev1.lower()}, dimostrano non solo la mia capacità di {theme1['keywords'][0]} ma anche un allineamento con l'approccio orientato al risultato che {company} incarna.

Inoltre, la mia esperienza in {theme2['name']} ({ev2.lower()}) completerebbe il vostro team attuale, contribuendo a guidare l'azienda verso nuovi traguardi.

La ringrazio per aver preso in considerazione la mia candidatura e mi rendo disponibile per un colloquio conoscitivo.

Cordiali Saluti,

Pierluigi Monaco
"""
    doc_cl.add_paragraph(body)
    
    buffer_cl = io.BytesIO()
    doc_cl.save(buffer_cl)
    buffer_cl.seek(0)

    # --- 2. CV OTTIMIZZATO (Formato ATS Lineare) ---
    doc_cv = Document()
    doc_cv.add_heading(meta['user_id'].replace('_',' ').title(), 0)
    doc_cv.add_paragraph(f"{meta['contact_info']['location']} | {meta['contact_info']['phone']} | {meta['contact_info']['email']}")
    
    doc_cv.add_heading('PROFESSIONAL SUMMARY' if lang == 'en' else 'PROFILO PROFESSIONALE', 1)
    doc_cv.add_paragraph(content['summary'])
    
    doc_cv.add_heading('CORE PROJECTS' if lang == 'en' else 'PROGETTI CHIAVE', 1)
    for proj in content['core_projects']:
        p = doc_cv.add_paragraph(style='List Bullet')
        p.add_run(f"{proj['title']} @ {proj['company']}: ").bold = True
        p.add_run(proj['impact'])
        
    doc_cv.add_heading('EXPERIENCE' if lang == 'en' else 'ESPERIENZA', 1)
    # Ruolo 1 (Tailored)
    p = doc_cv.add_paragraph()
    p.add_run("Strategic Finance Executive\n").bold = True
    p.add_run("OpenText / Micro Focus | 2018 - Present")
    doc_cv.add_paragraph(ev1, style='List Bullet') # Evidence mirata 1
    doc_cv.add_paragraph(ev2, style='List Bullet') # Evidence mirata 2
    
    # Ruolo 2 (Statico)
    p = doc_cv.add_paragraph()
    p.add_run("\nFinance Lead\n").bold = True
    p.add_run("Hewlett Packard | 2007 - 2018")
    doc_cv.add_paragraph("Managed P&L for Support & Consulting in EMEA." if lang == 'en' else "Gestione P&L per Support & Consulting in EMEA.", style='List Bullet')
    
    doc_cv.add_heading('SKILLS & EDUCATION' if lang == 'en' else 'COMPETENZE E FORMAZIONE', 1)
    doc_cv.add_paragraph(f"Hard Skills: {', '.join(content['skills']['hard'])}")
    doc_cv.add_paragraph(f"Tools: {', '.join(content['skills']['tools'])}")
    for edu in content['education']:
        doc_cv.add_paragraph(edu, style='List Bullet')

    buffer_cv = io.BytesIO()
    doc_cv.save(buffer_cv)
    buffer_cv.seek(0)
    
    return buffer_cl, buffer_cv

# --- INTERFACCIA STREAMLIT ---
st.title("🚀 Strategic Career Optimizer")
st.markdown("**Pierluigi Monaco** | Executive CV & Cover Letter Generator")
st.markdown("---")

col1, col2 = st.columns([1, 1])

with col1:
    st.markdown("### 1. Inserisci Annuncio")
    jd_input = st.text_area("Incolla qui la Job Description:", height=300, placeholder="Incolla il testo completo dell'annuncio...")

with col2:
    st.markdown("### 2. Risultati")
    if st.button("✨ Analizza e Genera Documenti", type="primary"):
        if len(jd_input) < 20:
            st.warning("⚠️ Il testo dell'annuncio è troppo breve.")
        else:
            # Carica dati
            profile = yaml.safe_load(PROFILE_YAML)
            
            # Analisi
            lang = detect_language(jd_input)
            company, role, char = extract_info(jd_input, lang)
            
            st.success(f"✅ **Analisi Completata!**")
            st.info(f"🌍 Lingua: **{lang.upper()}**\n🏢 Azienda: **{company}**\n🎯 Ruolo: **{role}**\n🔑 Key Value: **{char}**")
            
            # Generazione
            cl_buffer, cv_buffer = create_docs(profile, lang, company, role, char)
            
            st.markdown("### 📥 Download")
            st.download_button(
                label="📄 Scarica Cover Letter (.docx)",
                data=cl_buffer,
                file_name=f"Cover_Letter_{company}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            st.download_button(
                label="📄 Scarica CV Ottimizzato (.docx)",
                data=cv_buffer,
                file_name=f"CV_Optimized_{company}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            # Link Contatti
            st.markdown("---")
            st.markdown("### 🕵️ Headhunter Radar")
            q_hr = urllib.parse.quote(f'site:linkedin.com/in/ "{company}" "Talent Acquisition"')
            q_mgr = urllib.parse.quote(f'site:linkedin.com/in/ "{company}" "Finance Director" OR "CFO"')
            
            st.markdown(f"- [🔎 Cerca **HR / Talent Acquisition** su LinkedIn]({f'https://www.google.com/search?q={q_hr}'})")
            st.markdown(f"- [🔎 Cerca **Hiring Manager (CFO)** su LinkedIn]({f'https://www.google.com/search?q={q_mgr}'})")